from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
import sys
from Survival_analysis_feature_selection import Feature_selection_survival
from Survival_analysis_machine_learning import Data_analysis_survival
import copy
import pickle
import random
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
import win32com.client
import socket
import pandas as pd
from docx.shared import Inches
import numpy as np
from sklearn.linear_model import Lasso
import time
import os
import re
import shutil
import numpy as np
from sklearn.model_selection import train_test_split
import rpy2.robjects as robjects
import tempfile
import pickle
import winreg
def get_desktop():
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
    return(winreg.QueryValueEx(key, "Desktop")[0])
the_user_desktop = get_desktop()
dirname = tempfile.TemporaryDirectory()
class port(object):
    def __init__(self,view):
        self.view = view
    def write(self,*args):
        self.view.append(*args)
class Step_by_step_survival():
    def __init__(self):
        # initialize two objects in the class
        self.data = None
        self.model = None
        self.seed = random.randint(0, 10000)
        # define the operative times for "undo", "redo" and "clear"
        self.operative_times = 0
        self.operative_process_model_package = list()
        self.withdrawn_operative_process_model_package = list()
        self.operative_process = dict()
        self.withdrawn_operative_times_operative_process = dict()
        self.redo_times_for_reset = None
        self.model_load_path = None
        self.input_data = None
        self.standadization = None
        self.model_name = None
        ### define the functions
    def Select_a_path(self):
        #path = QtWidgets.QFileDialog.getOpenFileName(self, r"open file dialog", r"C:\Users\Administrator\Desktop", r"(*.csv)")
        path = QtWidgets.QFileDialog.getOpenFileName(self, r"open file dialog", the_user_desktop, r"(*.csv)")
        if (path[0]):
            self.operative_times += 1
            self.Result_text_survival.append("\n【Operative time: %s】" %self.operative_times)
            self.Result_text_survival.append("\n【Note】Choose a file path: %s \n" %path[0])
            self.Step1_data_path_survival.setText(str(path[0]))
            self.path = path[0]
            self.operative_process[self.operative_times] = ["select a path", [self.path]]
            global page1_data_path
            _, page1_data_path= os.path.split(self.path)
    def _Set_a_seed(self):
        self.operative_times += 1
        self.Result_text_survival.append("\n【Operative time: %s】" %self.operative_times)
        self.seed = int(self.Step1_seed_survival.text())
        self.Result_text_survival.append("\n【Note】Set a seed for data analysis: %s \n" %self.seed)
        self.operative_process[self.operative_times] = ["set a seed", [self.seed]]
    def Set_a_seed(self):
        seed = copy.copy(self.seed)
        try:
            self._Set_a_seed()
        except Exception as e:
            self.seed = copy.copy(seed)
            self.operative_times += -1
            self.Result_text_survival.append("\n【Error】Seed error: %s"%e)
    def _Seperate_a_data(self):
        self.operative_times += 1
        self.Result_text_survival.append("\n【Operative time: %s】" %self.operative_times)
        self.Result_text_survival.append("\n【Note】Seperate the data into two database: train and test ")
        self.Seperate_rate = float(self.Step1_seperate_survival.text())
        self.data = Feature_selection_survival(str(self.path))
        self.data.split(train_size=self.Seperate_rate, seed = self.seed)
        self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
        self.operative_process[self.operative_times] = ["seperate a data", [self.Seperate_rate, self.data.split_report_result_list]]
    def Seperate_a_data(self):
        try:
            self._Seperate_a_data()
        except Exception as e:
            self.operative_times += -1
            self.Result_text_survival.append("\n【Error】: %s"%e)
    def _Input_data(self):
        self.operative_times += 1
        self.Result_text_survival.append("\n【Operative time: %s】" %self.operative_times)
        if (str(self.Step2_inpute_survival.currentText()) == "Mean"):
            self.input_data = "Mean"
            self.Result_text_survival.append("\n【Note】Input data by 'Mean'")
            self.data.imputer1(is_abnormal_nan = self.Step2_inpute_is_abnormal_nan_survival.isChecked())
        elif (str(self.Step2_inpute_survival.currentText()) == "Median"):
            self.input_data = "Median"
            self.Result_text_survival.append("\n【Note】Input data by 'Median'")
            self.data.imputer2(is_abnormal_nan = self.Step2_inpute_is_abnormal_nan_survival.isChecked())
        self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
        self.operative_process[self.operative_times] = ["input data", [str(self.Step2_inpute_survival.currentText())]]
    def Input_data(self):
        data = copy.copy(self.data)
        try:
            self._Input_data()
        except Exception as e:
            self.operative_times += -1
            self.Result_text_survival.append("\n【Error】: %s"%e)
            self.data = copy.copy(data)
    def _Standarize_data(self):
        self.operative_times += 1
        self.Result_text_survival.append("\n【Operative time: %s】" %self.operative_times)
        if (str(self.Step2_standardize_survival.currentText()) == "Standardization"):
            self.standadization = "Standardization"
            self.Result_text_survival.append("\n【Note】Standarize data by 'Standardization'")
            self.data.Standard(method = "Standardization")
        elif (str(self.Step2_standardize_survival.currentText()) == "MinMaxScaler"):
            self.standadization = "MinMaxScaler"
            self.Result_text_survival.append("\n【Note】Standarize data by 'MinMaxScaler'")
            self.data.Standard(method = "MinMaxScaler")
        self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
        self.operative_process[self.operative_times] = ["standardize data", [str(self.Step2_standardize_survival.currentText())]]
    def Standarize_data(self):
        data = copy.copy(self.data)
        try:
            self._Standarize_data()
        except Exception as e:
            self.operative_times += -1
            self.Result_text_survival.append("\n【Error】: %s"%e)
            self.data = copy.copy(data)
    #单
    def _Select_feature1(self):
        self.operative_times += 1
        self.Result_text_survival.append("\n【Operative time: %s】" %self.operative_times)
        if (str(self.Step3_Univariate_method_survival.currentText()) == "Correlation_xx"):
            self.corr = float(self.Step3_Univariate_Correlation_corr_survival.text())
            self.Result_text_survival.append("\n【Note】Select feature by 'Correlation_xx'")
            self.Result_text_survival.append("\n【Note】Set corr: %.3f" %self.corr)
            self.data.select_by_correlation_xx(cutoff = self.corr)
            parameters = copy.copy(self.data.parameters)
            num_of_factors = len(self.data.data_variable_names)
            remained_features = copy.copy(self.data.data_variable_names)
            figure_report1 = copy.copy(self.data.figure_report1)
            try:
                figure_report2 = copy.copy(self.data.figure_report2)
            except:
                figure_report2 = None
            figure_report3 = copy.copy(self.data.figure_report3)
            try:
                figure_report4 = copy.copy(self.data.figure_report4)
            except:
                figure_report4 = None
            self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
            self.operative_process[self.operative_times] = ["select feature",
                                                            [str(self.Step3_Univariate_method_survival.currentText()), parameters,
                                                             num_of_factors, remained_features, figure_report1, figure_report2, figure_report3, figure_report4]]
        elif (str(self.Step3_Univariate_method_survival.currentText()) == "Univariate_Cox"):
            self.Univariate_Logistic_P_in = float(self.Step3_Univariate_Logistic_P_in_survival.text())
            self.Result_text_survival.append("\n【Note】Select feature by 'Univariate Cox'")
            self.Result_text_survival.append("\n【Note】Set corr: %.3f" % self.Univariate_Logistic_P_in)
            self.data.select_by_univariate_cox()
            parameters = copy.copy(self.data.parameters)
            num_of_factors = len(self.data.data_variable_names)
            remained_features = copy.copy(self.data.data_variable_names)
            figure_report1 = copy.copy(self.data.figure_report1)
            try:
                figure_report2 = copy.copy(self.data.figure_report2)
            except:
                figure_report2 = None
            figure_report3 = copy.copy(self.data.figure_report3)
            try:
                figure_report4 = copy.copy(self.data.figure_report4)
            except:
                figure_report4 = None
            self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
            self.operative_process[self.operative_times] = ["select feature",
                                                            [str(self.Step3_Univariate_method_survival.currentText()), parameters,
                                                             num_of_factors, remained_features, figure_report1, figure_report2, figure_report3, figure_report4]]
        elif (str(self.Step3_Univariate_method_survival.currentText()) == "Logrank_test"):
            self.Univariate_Logistic_P_in = float(self.Step3_Univariate_Logistic_P_in_survival.text())
            self.Result_text_survival.append("\n【Note】Select feature by 'Logrank test'")
            self.Result_text_survival.append("\n【Note】Set corr: %.3f" % self.Univariate_Logistic_P_in)
            self.data.select_by_logrank()
            parameters = copy.copy(self.data.parameters)
            num_of_factors = len(self.data.data_variable_names)
            remained_features = copy.copy(self.data.data_variable_names)
            figure_report1 = copy.copy(self.data.figure_report1)
            try:
                figure_report2 = copy.copy(self.data.figure_report2)
            except:
                figure_report2 = None
                
            figure_report3 = copy.copy(self.data.figure_report3)
            try:
                figure_report4 = copy.copy(self.data.figure_report4)
            except:
                figure_report4 = None
            self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
            self.operative_process[self.operative_times] = ["select feature",
                                                            [str(self.Step3_Univariate_method_survival.currentText()), parameters,
                                                             num_of_factors, remained_features, figure_report1, figure_report2, figure_report3, figure_report4]]
    def Select_feature1(self):
        data = copy.copy(self.data)
        data_variable_names = copy.copy(self.data.data_variable_names)
        operative_process = copy.copy(self.operative_process)
        try:
            self._Select_feature1()
            plt.show()
            self.Result_text_survival.append("Successfully Done")
        except Exception as e:
            self.operative_times += -1
            self.Result_text_survival.append("\n【Error】: %s " %e)
            self.data = copy.copy(data)
            self.data.data_variable_names = copy.copy(data_variable_names)
            self.operative_process = copy.copy(operative_process)
    def _Select_feature2(self):
        self.operative_times += 1
        self.Result_text_survival.append("\n【Operative time: %s】" % self.operative_times)
        if (str(self.Step3_MultiVariate_method_survival.currentText()) == "MultiVariate_Cox"):
            if (str(self.Step3_MultiVariate_Logistic_Result_survival.currentText()) == "Features"):
                self.select = self.Step3_MultiVariate_Logistic_Features_Select_survival.isChecked()
                self.Result_text_survival.append("\n【Note】Select feature by 'MultiVariate Cox'")
                self.data.select_by_stepwise_cox(select = self.select, radscore = False)
                parameters = copy.copy(self.data.parameters)
                num_of_factors = len(self.data.data_variable_names)
                remained_features = copy.copy(self.data.data_variable_names)
                figure_report1 = copy.copy(self.data.figure_report1)
                try:
                    figure_report2 = copy.copy(self.data.figure_report2)
                except:
                    figure_report2 = None
                figure_report3 = copy.copy(self.data.figure_report3)
                try:
                    figure_report4 = copy.copy(self.data.figure_report4)
                except:
                    figure_report4 = None
                self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
                self.operative_process[self.operative_times] = ["select feature",
                                                                [str(self.Step3_MultiVariate_method_survival.currentText()), parameters,
                                                                 num_of_factors, remained_features, figure_report1, figure_report2, figure_report3, figure_report4]]
            elif (str(self.Step3_MultiVariate_Logistic_Result_survival.currentText()) == "Radscore"):
                self.select = self.Step3_MultiVariate_Logistic_Radscore_Select_survival.isChecked()
                print(self.select)
                self.Result_text_survival.append("\n【Note】Show Rad_score")
                self.data.select_by_stepwise_cox(select = self.select, radscore = True)
                parameters = copy.copy(self.data.parameters)
                num_of_factors = len(self.data.data_variable_names)
                remained_features = copy.copy(self.data.data_variable_names)
                self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
                self.operative_process[self.operative_times] = ["select feature",
                                                                [str(self.Step3_MultiVariate_method_survival.currentText()), parameters,
                                                                 num_of_factors, remained_features]]
        elif (str(self.Step3_MultiVariate_method_survival.currentText()) == "L1"):
            self.Result_text_survival.append("\n【Note】Select feature by 'L1'")
            self.data.select_by_lasso_cox()
            parameters = copy.copy(self.data.parameters)
            num_of_factors = len(self.data.data_variable_names)
            remained_features = copy.copy(self.data.data_variable_names)
            figure_report1 = copy.copy(self.data.figure_report1)
            try:
                figure_report2 = copy.copy(self.data.figure_report2)
            except:
                figure_report2 = None
            figure_report3 = copy.copy(self.data.figure_report3)
            try:
                figure_report4 = copy.copy(self.data.figure_report4)
            except:
                figure_report4 = None
            figure_report5 = copy.copy(self.data.lasso_figure_1)
            figure_report6 = copy.copy(self.data.lasso_figure_2)
            self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
            self.operative_process[self.operative_times] = ["select feature",
                                                            [str(self.Step3_MultiVariate_method_survival.currentText()), parameters,
                                                             num_of_factors, remained_features, figure_report1, figure_report2, figure_report3, figure_report4, figure_report5, figure_report6]]
    def Select_feature2(self):
        data = copy.copy(self.data)
        data_variable_names = copy.copy(self.data.data_variable_names)
        operative_process = copy.copy(self.operative_process)
        try:
            self._Select_feature2()
            plt.show()
            self.Result_text_survival.append("Successfully Done")
        except Exception as e:
            self.operative_times += -1
            self.Result_text_survival.append("\n【Error】: %s " % e)
            self.data = copy.copy(data)
            self.data.data_variable_names = copy.copy(data_variable_names)
            self.operative_process = copy.copy(operative_process)
    def _Machine_learning(self):
        self.operative_times += 1
        self.Result_text_survival.append("\n【Operative time: %s】" %self.operative_times)
        if self.data.test:
            whether_test = True
            self.model = Data_analysis_survival(self.data.train_feature, self.data.train_time_status, self.data.test_feature, self.data.test_time_status, self.data.data_variable_names, whether_test = whether_test, seed = self.seed, cross = self.Step4_method_Cross_survival.isChecked())
        else:
            whether_test = False
            self.model = Data_analysis_survival(self.data.train_feature, self.data.train_time_status, None, None, data_variable_names = self.data.data_variable_names, whether_test = whether_test, seed = self.seed, cross = self.Step4_method_Cross_survival.isChecked())
        if (str(self.Step4_method_survival.currentText()) == "Cox_PH_model"):
            self.Result_text_survival.append("\n【Note】Use 'Cox proportional hazard' model")
            self.model_name = str(self.Step4_method_survival.currentText())
            if int(self.Step4_Method_Logistic_time_point2_survival.text()) == -1:
                time_point2 = False
            else:
                time_point2 = int(self.Step4_Method_Logistic_time_point2_survival.text())
            self.model.multivariate_analysis(Nomo = self.Step4_method_Auto_or_Nomo_survival.isChecked(), time_point1 = int(self.Step4_Method_Logistic_time_point1_survival.text()), time_point2 = time_point2, figure_label1 = "training figure", figure_label2 = "testing figure")
            statistic_information_summary = copy.copy(self.model.train_multivariate_result)
            if whether_test == True:
                report_result_list = copy.copy([self.model.train_c_index, self.model.test_c_index])
            else:
                report_result_list = copy.copy([self.model.train_c_index])
            figure_report1 = self.model.tdROC_figure1
            figure_report2 = self.model.tdROC_figure2
            figure_report3 = self.model.calibrate_figure1
            figure_report4 = self.model.calibrate_figure2
            Nomogram_figure = self.model.Nomogram_figure
            if self.Step4_method_Cross_survival.isChecked():
                self.operative_process[self.operative_times] = ["machine_learning",[str(self.Step4_method_survival.currentText()), statistic_information_summary, report_result_list, figure_report1, figure_report2, figure_report3, figure_report4, Nomogram_figure, list(self.model.cross_validation_value.T[1]), list(self.model.cross_validation_mean)[1]]]
            else:
                self.operative_process[self.operative_times] = ["machine_learning",[str(self.Step4_method_survival.currentText()), statistic_information_summary, report_result_list, figure_report1, figure_report2, figure_report3, figure_report4, Nomogram_figure, None, None]]
            self.operative_process_model_package.append((copy.copy(self.data), copy.copy(self.model),copy.copy(self.input_data),copy.copy(self.standadization)))
    def Machine_learning(self):
        model = copy.copy(self.model)
        operative_process = copy.copy(self.operative_process)
        try:    
            self._Machine_learning()
            plt.show()
        except Exception as e:
            self.operative_times += -1
            self.Result_text_survival.append("\n【Error】: %s" %e)
            self.model = copy.copy(model)
            self.operative_process = copy.copy(operative_process)
    def Data_select_a_path(self):
        path = QtWidgets.QFileDialog.getSaveFileName(self, r"Bulid a file", the_user_desktop,r"(*.csv)")
        if (path[0]):
            self.Result_text_survival.append("\n【Note】Assign a filename for saving data %s" % path[0])
            self.Save_Data_path_survival.setText(str(path[0]))
            self.Data_save_path = path[0]
    def _Data_save(self):
        if self.data.test:
            train_sample_ones = np.ones(self.data.train_feature.shape[0])
            test_sample_zeros = np.zeros(self.data.test_feature.shape[0])
            sample_ones_zeros = np.hstack((train_sample_ones, test_sample_zeros))
            feature = np.vstack((self.data.train_feature, self.data.test_feature))
            label = np.vstack((self.data.train_time_status, self.data.test_time_status))
            data_index = np.arange(0, feature.shape[0])
            data_index = data_index.reshape(len(data_index), 1)
            train_index, test_index = train_test_split(data_index, train_size = self.Seperate_rate, random_state = self.seed)
            index_after_split = np.vstack((train_index, test_index))
            sample_ones_zeros = sample_ones_zeros.reshape(sample_ones_zeros.shape[0], 1)
            data = np.hstack((index_after_split, label, sample_ones_zeros, feature))
            data = pd.DataFrame(data, columns = ["index", "times", "status", "whether_train"] + self.data.data_variable_names)
            data.sort_values(by = "index", ascending = True, inplace = True)
            data.to_csv(self.Data_save_path, index = False)
        else:
            data = np.hstack((self.data.train_time_status, self.data.train_feature))
            data = pd.DataFrame(data, columns = ["times", "status"] + self.data.data_variable_names)
            data.to_csv(self.Data_save_path, index = False)
        self.Result_text_survival.append("\n【Note】Save the data in %s" % self.Data_save_path)
    def Data_save(self):
        try:
            self._Data_save()
        except Exception as e:
            self.Result_text_survival.append("\n【Error】%s"%e)
    def Result_select_a_path(self):
        path = QtWidgets.QFileDialog.getSaveFileName(self, r"Bulid a file", the_user_desktop, r"(*.doc)")
        if (path[0]):
            self.Result_text_survival.append("\n【Note】Assign a filename for saving results %s" %path[0])
            self.Save_result_path_survival.setText(str(path[0]))
            self.Result_save_path = path[0]
    def _Result_save(self):
        self.Whole_report()
        self.Result_text_survival.append("\n【Note】Save the results in %s" %self.Result_save_path)
        the_desktop_path = get_desktop()
        shutil.copyfile(os.path.dirname(os.path.abspath(__file__)) + "\\description_template\\cox_method_template.docx", the_desktop_path + "\\cox_method_template.docx")
        self.Result_text_survival.append("\n【Note】The step_by_step method's template is exported to the desktop.")
    def Result_save(self):
        try:
            self._Result_save()
        except Exception as e:
            self.Result_text_survival.append("\n【Error】%s"%e)
    # define the "undo", "redo" and "clear"
    def Undo_operation(self):
        self.operative_times += -1
        if (self.operative_times == 0 or self.operative_process_model_package == []):
            self.operative_times += 1
            self.Result_text_survival.append("\n【Reach the initial operation】 ")
            pass
        else:
            self.Result_text_survival.append("\n【Come back to the operative time %s】 %s:%s" %(self.operative_times,
                                                                                self.operative_process[self.operative_times][0],
                                                                                self.operative_process[self.operative_times][1][0]))

            self.withdrawn_operative_process_model_package.append(self.operative_process_model_package.pop())
            self.withdrawn_operative_times_operative_process[self.operative_times + 1] = (self.operative_process.pop(self.operative_times + 1))
            if (len(self.operative_process_model_package) != 0):
                undo_data, undo_model, undo_input, undo_standardization = self.operative_process_model_package[len(self.operative_process_model_package) - 1]
                self.data = copy.copy(undo_data)
                self.model = copy.copy(undo_model)
                self.input_data = copy.copy(undo_input)
                self.standadization = copy.copy(undo_standardization)
            else:
                self.data = None
                self.model = None
                self.input_data = None
                self.standadization = None
            # define the times in order to reset the self.withdrawn_operative_process_model_package
            self.redo_times_for_reset = len(self.operative_process_model_package)
    def Redo_operation(self):
        self.operative_times += 1
        # define when to reset the self.withdrawn_operative_process_model_package
        if self.redo_times_for_reset != len(self.operative_process_model_package):
            self.operative_times += -1
            self.withdrawn_operative_process_model_package.clear()
            self.withdrawn_operative_times_operative_process.clear()
            self.Result_text_survival.append("\n【Already reach the last operation】 ")
            pass
        elif (self.withdrawn_operative_process_model_package == []):
        # define the redo process
            self.operative_times += -1
            self.Result_text_survival.append("\n【Already reach the last operation】 ")
            pass
        else:
            self.Result_text_survival.append("\n【Go forward to the operative time %s】 %s:%s" %(self.operative_times,
                                                              self.withdrawn_operative_times_operative_process[self.operative_times][0],
                                                              self.withdrawn_operative_times_operative_process[self.operative_times][1][0]))
            
            self.operative_process[self.operative_times] = self.withdrawn_operative_times_operative_process[self.operative_times]
            self.operative_process_model_package.append(self.withdrawn_operative_process_model_package.pop())
            self.withdrawn_operative_times_operative_process[self.operative_times] = self.withdrawn_operative_times_operative_process.pop(self.operative_times)
            redo_data, redo_model, redo_input, redo_standardization = self.operative_process_model_package[len(self.operative_process_model_package) - 1]
            self.data = copy.copy(redo_data)
            self.model = copy.copy(redo_model)
            self.input_data = copy.copy(redo_input)
            self.standadization = copy.copy(redo_standardization)
    def Model_select_a_path(self):
        path = QtWidgets.QFileDialog.getSaveFileName(self, r"Bulid a file", the_user_desktop, r"(*.pkl)")
        if (path[0]):
            self.Result_text_survival.append("\n【Note】Assign a filename for saving model %s" %path[0])
            self.Save_model_path_survival.setText(str(path[0]))
            self.Model_save_path = path[0]
    def _Model_save(self):
        if self.model:
            train_data = np.concatenate((self.data.train_time_status, self.data.train_feature), axis = 1)
            train_data = pd.DataFrame(train_data, columns = ["time", "status"] + list(self.data.data_variable_names))
            model_txt = open(self.Model_save_path, "wb")
            pickle.dump([self.data, self.model, [self.input_data, self.standadization], train_data], model_txt)
            model_txt.close()
        elif self.data.whether_radscore:
            model_txt = open(self.Model_save_path, "wb")
            pickle.dump([self.data, self.model, [self.input_data, self.standadization]], model_txt)
            model_txt.close()
    def Model_save(self):
        try:
            self._Model_save()
            self.Result_text_survival.append("\n【Note】 Save successfully")
        except Exception as e:
            self.Result_text_survival.append("\n【Error】 %s"%e)
    def Load_model_select_a_path(self):
        path = QtWidgets.QFileDialog.getOpenFileName(self, r"open file dialog", the_user_desktop, r"(*.pkl)")
        if (path[0]):
            self.Result_text_survival.append("\n【Note】Choose a file path: %s" %path[0])
            self.load_model_path_survival.setText(str(path[0]))
            self.model_load_path = path[0]
    def Model_load(self):
        if self.model_load_path:
            self.Predict_module_predict = False
            self.Predict_module_radscore = False
          
            try:
                self.labeled_trained_model_data = pickle.load(open(self.model_load_path, 'rb'))[0] #Load the data of the predictive model
                model = pickle.load(open(self.model_load_path, 'rb'))[1] #载入的模型类
                self.preprocessing = pickle.load(open(self.model_load_path, 'rb'))[2] #The process of data set preprocessing, that is, what kind of missing values are filled and what kind of data is standardized
                self.train_data_for_radscore = pickle.load(open(self.model_load_path, 'rb'))[3] 
                self.features = model.data_variable_names #Feature name used in the model
                print(self.train_data_for_radscore)
                self.Predict_module_predict = True
                self.Result_text_survival.append("\n【Note】load the chosen model object file.")
                print("【Note】The predict model is loaded.")
            except:
                self.labeled_trained_model_data = pickle.load(open(self.model_load_path, 'rb'))[0]
                self.preprocessing = pickle.load(open(self.model_load_path, 'rb'))[2]
                self.data_variable_names_for_radscore = self.labeled_trained_model_data.data_variable_names_for_radscore
                self.coef_after_radscore = self.labeled_trained_model_data.coef_after_radscore
                self.Predict_module_radscore = True
                self.Result_text_survival.append("\n【Note】load the chosen model object file.")
                print("【Note】The radscore model is loaded.")
        else:
            self.Result_text_survival.append("\n There is no model object in system, neither one been loaded.")
    def Predict_data_load_select_a_path(self):
        path = QtWidgets.QFileDialog.getOpenFileName(self, r"open file dialog", the_user_desktop, r"(*.csv)")
        if (path[0]):
            self.Result_text_survival.append("\n【Note】Data file been chosen in path: %s" %path[0])
            self.Load_data_path_survival.setText(str(path[0]))
            self.Data_predict_path = path[0]
    def data_predict_(self):
        self.labeled_data = pd.read_csv(self.Data_predict_path,engine = 'python') #Read in the data that needs to be predicted
        # If it is used to predict the outcome
        if self.Predict_module_predict:
            data = self.labeled_data.values[:,2:]
            if self.preprocessing[0]:
                data = self.labeled_trained_model_data.imp.transform(data)
            if self.preprocessing[1]:
                data = self.labeled_trained_model_data.scaler.transform(data)
            self.labeled_data_after_transform = pd.DataFrame(data, columns=self.labeled_data.columns[2:])   # The data has been preprocessed, now it is pd.DataFrame
            self.labeled_data_after_transform = self.labeled_data_after_transform[self.features]
            predict_data_path = "\\".join([dirname.name, "".join([str(time.time()), "predict_data_path.csv"])])
            self.labeled_data_after_transform.to_csv(predict_data_path, index = False)
            train_data_path = "\\".join([dirname.name, "".join([str(time.time()), "train_data_path.csv"])])
            self.train_data_for_radscore.to_csv(train_data_path, index = False)
            export_data_path = "\\".join([dirname.name, "".join([str(time.time()), "export_data_path.csv"])])
            r = robjects.r
            r('''
            library(survival)
            ''')
            r('''
            import_to_r <- function(train_data_path_r, predict_data_path_r, export_data_path_r){
                train_data_path <<- train_data_path_r
                predict_data_path <<- predict_data_path_r
                export_data_path <<- export_data_path_r
                }
            ''')
            r['import_to_r'](train_data_path, predict_data_path, export_data_path)
            r('''
            data <- read.csv(train_data_path, header = TRUE, sep = ",")
            f <- coxph(Surv(time, status) ~ ., data)           
            predict_data <- read.csv(predict_data_path, header = TRUE, sep = ",")
            result <- predict(f, predict_data)           
            write.csv(result, export_data_path)
            ''')
            results = pd.read_csv(export_data_path)
            results = results.values[:,1]
            print("\n【Note】predicting labels:\n %s" %results)
            self.labeled_data_after_transform.insert(0, 'Predict', results)
       
        if self.Predict_module_radscore:
            data = self.labeled_data.values[:,0:]
            if self.preprocessing[0]:
                data = self.labeled_trained_model_data.imp.transform(data)
            if self.preprocessing[1]:
                data = self.labeled_trained_model_data.scaler.transform(data)
            self.labeled_data_after_transform = pd.DataFrame(data, columns=self.labeled_data.columns[0:])
            self.labeled_data_after_transform = self.labeled_data_after_transform[self.data_variable_names_for_radscore]
            radscore = np.dot(self.coef_after_radscore, self.labeled_data_after_transform.T)
            self.labeled_data_after_transform.insert(0, 'rad_score', radscore)
            print("\n【Note】The radscore is calculated.")
    def data_predict(self):
        try:
            self.data_predict_()
        except Exception as e:
            self.Result_text.append("\n【Error】 %s" %e)
    def Predict_data_save_select_a_path(self):
        path = QtWidgets.QFileDialog.getSaveFileName(self, r"Bulid a file", the_user_desktop,r"(*.csv)")
        if (path[0]):
            self.Result_text_survival.append("\n【Note】Assign a filename for saving data %s" % path[0])
            self.Predict_Save_Data_path_survival.setText(str(path[0]))
            self.Predict_data_save_path = path[0]
    def Predict_data_save(self):
        try:
            self.labeled_data_after_transform.to_csv(self.Predict_data_save_path,index = False)
            self.Result_text_survival.append("\n【Note】Save the data in %s" % self.Predict_data_save_path)
        except Exception as e:
            self.Result_text.append("\n【Error】 %s"%e)
    def Clear_operation(self):
        # self.operative_times = 0
        # self.operative_process_model_package.clear()
        # self.withdrawn_operative_process_model_package.clear()
        self.operative_process.clear()
        self.withdrawn_operative_times_operative_process.clear()
        self.data, self.model = None, None
        self.model_name = None
        self.input_data = None
        self.standadization = None
        self.Step1_data_path_survival.clear()
        self.Step1_seed_survival.clear()
        self.Step1_seperate_survival.setValue(0.7)
        self.Step2_inpute_survival.setCurrentText("Median")
        self.Step2_standardize_survival.setCurrentText("Standardization")
        self.Step3_Factors_survival.setCurrentIndex(0)
        self.Step3_Univariate_method_survival.setCurrentIndex(0)
        # self.Step4_method_survival.setCurrentText("Cox_PH_model")
        self.Step2_inpute_is_abnormal_nan_survival.setChecked(False)
        # self.Save_model_path_survival.clear()
        # self.Save_result_path_survival.clear()
        self.Save_Data_path_survival.clear()
        # self.load_model_path_survival.clear()
        # self.Load_data_path_survival.clear()
        # self.Predict_Save_Data_path_survival.clear()
        self.path = None
        # self.seed = None
        self.Result_text_survival.clear()
        self.Result_text_survival.append("\n【Clear all operations, operation step is initialized to 0】 \n")
    def Whole_report(self):
        word = Document()
        ### 增加封面 ###
        word.add_paragraph("")
        word.add_paragraph("")
        word.add_paragraph("")
        word.add_paragraph("")
        word.add_paragraph("")
        word.add_paragraph("")
        paragraph_page1_1 = word.add_paragraph()
        paragraph_page1_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_run1_1 = paragraph_page1_1.add_run("Statistical Report")
        paragraph_run1_1.bold = True
        paragraph_run1_1.font.color.rgb = RGBColor(17, 30, 95)
        paragraph_run1_1.font.size = Pt(42)
        word.add_paragraph("")
        word.add_paragraph("")
        word.add_paragraph("")
        paragraph_page1_3 = word.add_paragraph()
        paragraph_run1_3 = paragraph_page1_3.add_run("\t\tAuthor: %s" %socket.gethostname())
        paragraph_run1_3.font.color.rgb = RGBColor(17, 30, 95)
        paragraph_run1_3.font.size = Pt(18)
        paragraph_page1_4 = word.add_paragraph()
        paragraph_run1_4 = paragraph_page1_4.add_run("\t\tTime: %s" %time.strftime('%Y-%m-%d %H:%M:%S',time.localtime(time.time())))
        paragraph_run1_4.font.color.rgb = RGBColor(17, 30, 95)
        paragraph_run1_4.font.size = Pt(18)
        paragraph_page1_5 = word.add_paragraph()
        paragraph_run1_5 = paragraph_page1_5.add_run("\t\tData: '%s'" %page1_data_path)
        paragraph_run1_5.font.color.rgb = RGBColor(17, 30, 95)
        paragraph_run1_5.font.size = Pt(18)
        word.add_page_break()
        ### The content of the official output report ###
        word.add_heading("Summary Report", 0)
        # Obtain the best machine learning model and locate its location
        is_machine_learning = False
        machine_learning_list_step = []
        machine_learning_list_name = []
        machine_learning_list_c_index = []
        machine_learning_step = 1
        for i in np.arange(1, self.operative_times + 1):
            if self.operative_process[i][0] == "machine_learning":
                # Extract machine learning steps, names, accuracy, AUC; machine learning steps are easy to track the previous steps of the best model
                machine_learning_list_step.append(i)
                machine_learning_list_name.append(self.operative_process[i][0]+str(machine_learning_step))
                machine_learning_list_c_index.append(self.operative_process[i][1][2][0]['c-index'])
                machine_learning_step+=1
                is_machine_learning = True
            else:
                pass
        if is_machine_learning:
            machine_learning_list_c_index = np.array(machine_learning_list_c_index)
            # Output the position corresponding to the best accuracy
            where_train_index_max_index = np.array(np.where(machine_learning_list_c_index ==np.max(machine_learning_list_c_index)))
        else:
            pass
        # If machine learning has been performed, the feature selection step before the optimal model is obtained; if machine learning has not been performed, no summary report will be output
        if is_machine_learning:
            best_model_step = machine_learning_list_step[where_train_index_max_index[0][0]]
            training_process = []
            for i in np.arange(1, best_model_step):
                if self.operative_process[i][0] == "machine_learning":
                    pass
                else:
                    training_process.append(self.operative_process[i][0])
            training_process.append("machine_learning")
            word.add_heading("The summary report recognized the best model from all models your training, and summarized its training process.", 1)
            word.add_paragraph("As the result shown, the best model was %s, its training process including: %s. Detailed information is shown below:" %(self.operative_process[best_model_step][1][0], training_process))
            for i in np.arange(1, best_model_step):
                if (self.operative_process[i][0] == "select a path"):
                    text2 = "%s. Data: %s" %(i, self.operative_process[i][1][0])
                    word.add_heading(text2, 1)
                elif (self.operative_process[i][0] == "set a seed"):
                    text2 = "%s. Random seed: %s" %(i, self.operative_process[i][1][0])
                    word.add_heading(text2, 1)
                elif (self.operative_process[i][0] == "seperate a data"):
                    if self.data.test:
                        text2 = "%s. Seperative rate: %s" %(i, self.operative_process[i][1][0])
                        word.add_heading(text2, 1)
                elif (self.operative_process[i][0] == "input data"):
                    text2 = "%s. The method for filling the missing data: %s" %(i, self.operative_process[i][1][0])
                    word.add_heading(text2, 1)
                elif (self.operative_process[i][0] == "standardize data"):
                    text2 = "%s. The method for standardizing the data: %s" %(i, self.operative_process[i][1][0])
                    word.add_heading(text2, 1)
                elif (self.operative_process[i][0] == "select feature"):
                    for j in np.arange(0, len(self.operative_process[i][1])):
                        if (j == 3):
                            text2 = "%s. The method for selecting features: %s" %(i, self.operative_process[i][1][0])
                            word.add_heading(text2, 1)
                            if (self.operative_process[i][1][0] == "PCA"):
                                text3 = "parameters setted: %s \nnum of principal factors: %s \nThe expalined variance by principal factors: \n%s" %(self.operative_process[i][1][1], self.operative_process[i][1][2], self.operative_process[i][1][3])
                                word.add_paragraph(text3)
                            else:
                                text3 = "parameters setted: %s \nnum of remained features: %s \nremained features: \n%s" %(self.operative_process[i][1][1], self.operative_process[i][1][2], self.operative_process[i][1][3])
                                word.add_paragraph(text3)
                        elif (j == 4):
                            word.add_paragraph("Heatmap of the model in the training samples:")
                            word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                        elif (j == 5):
                            try:
                                word.add_paragraph("Heatmap of the model in the testing samples:")
                                word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                            except:
                                word.add_paragraph('')
                        elif (j == 6):
                            try:
                                word.add_paragraph("Correlation coefficient figure of the training samples")
                                word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                            except:
                                word.add_paragraph('')
                        elif (j == 7):
                            try:
                                word.add_paragraph("Correlation coefficient figure of the testing samples")
                                word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                            except:
                                word.add_paragraph('')
                        elif (j == 8):
                            word.add_paragraph("lasso path plot of the model in the training samples:")
                            word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                        elif (j == 9):
                            word.add_paragraph("Mean square error on each fold for lasso model:")
                            word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                elif (self.operative_process[i][0] == "machine_learning"):
                    pass
            i += 1
            text2 = "%s. The best machine learning method: %s" %(i, self.operative_process[i][1][0])
            word.add_heading(text2, 1)
            i = best_model_step
            if self.operative_process[i][1][0] == "Cox_PH_model":
                for j in np.arange(0, len(self.operative_process[i][1])):
                    if (j == 1):
                        word.add_paragraph("\nThe statistic information of Logistic models:")
                        statistic_table = self.operative_process[i][1][j]
                        statistic_table_value = statistic_table.values
                        report_table2 = word.add_table(rows = statistic_table_value.shape[0]+1, cols = 6, style = "Table Grid")
                        report_table2.autofit = False
                        report_table2.columns[0].width = Inches(0.1)
                        report_table2.cell(0,0).text = ['variables']
                        report_table2.cell(0,1).text = ['coef']
                        report_table2.cell(0,2).text = ['HR']
                        report_table2.cell(0,3).text = ['95% CI of coef']
                        report_table2.cell(0,4).text = ['z']
                        report_table2.cell(0,5).text = ['P']
                        length = statistic_table_value.shape[0]+1
                        for b in range(1, length):
                            report_table2.cell(b,0).text = str(statistic_table_value[b-1, 0])
                            report_table2.cell(b,1).text = str(round(statistic_table_value[b-1, 1], 3))
                            report_table2.cell(b,2).text = str(round(statistic_table_value[b-1, 2], 3))
                            report_table2.cell(b,3).text = str(round(statistic_table_value[b-1, 3], 3))
                            report_table2.cell(b,4).text = str(round(statistic_table_value[b-1, 4], 3))
                            report_table2.cell(b,5).text = str(round(statistic_table_value[b-1, 5], 3))
                    elif (j == 2):
                        word.add_paragraph('')
                        word.add_paragraph('c-index of the model:')
                        word.add_paragraph(str(self.operative_process[i][1][j]))
                    elif (j == 3):
                        word.add_paragraph('')
                        word.add_paragraph("tdROC of logistic model in the training samples:")
                        word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                    elif (j == 4):
                        try:
                            word.add_paragraph('')
                            word.add_paragraph("tdROC of logistic model in the testing samples:")
                            word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                        except:
                            word.add_paragraph('')
                # ======================================================Add nomo diagram and calibration diagram to the report==========================================
                    elif (j == 5):
                        word.add_paragraph("Calibration plot of logistic model in training samples:")
                        word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                    elif (j == 6):
                        try:
                            word.add_paragraph("Calibration plot of logistic model in testing samples:")
                            word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                        except:
                            word.add_paragraph('')
                    elif (j == 7):
                        try:
                            word.add_paragraph("Nomograme plot of the Logistic model in the training samples:")
                            word.add_picture(self.operative_process[i][1][j], height=Inches(6))
                        except:
                            word.add_paragraph("You had not made a nomogram")
                            
                    elif (j == 8):
                                word.add_paragraph("each-validation c-index:")
                                word.add_paragraph(str(self.operative_process[i][1][j]))
                    elif (j == 9):
                        word.add_paragraph("mean-validation c-index:")
                        word.add_paragraph(str(self.operative_process[i][1][j]))
        else:
            word.add_paragraph("No summary report: there is no machine learning model.")
        ### The content of the official output report ###
        if len(machine_learning_list_step) > 1:
            word.add_page_break()
            word.add_heading("Step by Step Report", 0)
            word.add_heading("The step by step report summarized all training processes your have done.", 1)
            for i in np.arange(1, self.operative_times + 1):
                text1 = "Operative time: %i, %s" %(i, self.operative_process[i][0])
                word.add_heading(text1, 1)
                if (self.operative_process[i][0] == "select a path"):
                    text2 = "you select the data file in %s" %self.operative_process[i][1][0]
                    word.add_paragraph(text2)
                elif (self.operative_process[i][0] == "set a seed"):
                    text2 = "you set the seed: %s, it will be used in all processes where seed is needed" %self.operative_process[i][1][0]
                    word.add_paragraph(text2)
                elif (self.operative_process[i][0] == "seperate a data"):
                    if self.data.test:
                        text2 = "%s. Seperative rate: %s" %(i, self.operative_process[i][1][0])
                        word.add_heading(text2, 1)
                elif (self.operative_process[i][0] == "input data"):
                    text2 = "the method for filling the missing data: %s" %self.operative_process[i][1][0]
                    word.add_paragraph(text2)
                elif (self.operative_process[i][0] == "standardize data"):
                    text2 = "the method for standardizing the data: %s" %self.operative_process[i][1][0]
                    word.add_paragraph(text2)
                elif (self.operative_process[i][0] == "select feature"):
                    for j in np.arange(0, len(self.operative_process[i][1])):
                        if (j == 3):
                            text2 = "%s. The method for selecting features: %s" %(i, self.operative_process[i][1][0])
                            word.add_heading(text2, 1)
                            if (self.operative_process[i][1][0] == "PCA"):
                                text3 = "parameters setted: %s \nnum of principal factors: %s \nThe expalined variance by principal factors: \n%s" %(self.operative_process[i][1][1], self.operative_process[i][1][2], self.operative_process[i][1][3])
                                word.add_paragraph(text3)
                            else:
                                text3 = "parameters setted: %s \nnum of remained features: %s \nremained features: \n%s" %(self.operative_process[i][1][1], self.operative_process[i][1][2], self.operative_process[i][1][3])
                                word.add_paragraph(text3)
                        elif (j == 4):
                            word.add_paragraph("Heatmap of the model in the training samples:")
                            word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                        elif (j == 5):
                            try:
                                word.add_paragraph("Heatmap of the model in the testing samples:")
                                word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                            except:
                                word.add_paragraph('')
                        elif (j == 6):
                            try:
                                word.add_paragraph("Correlation coefficient figure of the training samples")
                                word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                            except:
                                word.add_paragraph('')
                        elif (j == 7):
                            try:
                                word.add_paragraph("Correlation coefficient figure of the testing samples")
                                word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                            except:
                                word.add_paragraph('')
                        elif (j == 8):
                            word.add_paragraph("lasso path plot of the model in the training samples:")
                            word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                        elif (j == 9):
                            word.add_paragraph("Mean square error on each fold for lasso model:")
                            word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                elif (self.operative_process[i][0] == "machine_learning"):
                    if self.operative_process[i][1][0] == "Cox_PH_model":
                        for j in np.arange(0, len(self.operative_process[i][1])):
                            if (j == 1):
                                word.add_paragraph("\nThe statistic information of Logistic models:")
                                statistic_table = self.operative_process[i][1][j]
                                statistic_table_value = statistic_table.values
                                report_table2 = word.add_table(rows = statistic_table_value.shape[0]+1, cols = 6, style = "Table Grid")
                                report_table2.autofit = False
                                report_table2.columns[0].width = Inches(0.1)
                                report_table2.cell(0,0).text = ['variables']
                                report_table2.cell(0,1).text = ['coef']
                                report_table2.cell(0,2).text = ['HR']
                                report_table2.cell(0,3).text = ['95% CI of coef']
                                report_table2.cell(0,4).text = ['z']
                                report_table2.cell(0,5).text = ['P']
                                length = statistic_table_value.shape[0]+1
                                for b in range(1, length):
                                    report_table2.cell(b,0).text = str(statistic_table_value[b-1, 0])
                                    report_table2.cell(b,1).text = str(round(statistic_table_value[b-1, 1], 3))
                                    report_table2.cell(b,2).text = str(round(statistic_table_value[b-1, 2], 3))
                                    report_table2.cell(b,3).text = str(round(statistic_table_value[b-1, 3], 3))
                                    report_table2.cell(b,4).text = str(round(statistic_table_value[b-1, 4], 3))
                                    report_table2.cell(b,5).text = str(round(statistic_table_value[b-1, 5], 3))
                            elif (j == 2):
                                word.add_paragraph('')
                                word.add_paragraph('c-index of the model (training data & testing data):')
                                word.add_paragraph(str(self.operative_process[i][1][j]))
                            elif (j == 3):
                                word.add_paragraph('')
                                word.add_paragraph("tdROC of Cox model in the training samples:")
                                word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                            elif (j == 4):
                                try:
                                    word.add_paragraph('')
                                    word.add_paragraph("tdROC of Cox model in the testing samples:")
                                    word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                                except:
                                    word.add_paragraph('')
                        # ======================================================Add nomo diagram and calibration diagram to the report=========================================
                            elif (j == 5):
                                word.add_paragraph("Calibration plot of Cox model in training samples:")
                                word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                            elif (j == 6):
                                try:
                                    word.add_paragraph("Calibration plot of Cox model in testing samples:")
                                    word.add_picture(self.operative_process[i][1][j], height=Inches(4))
                                except:
                                    word.add_paragraph('')
                            elif (j == 7):
                                try:
                                    word.add_paragraph("Nomograme plot of the Cox model in the training samples:")
                                    word.add_picture(self.operative_process[i][1][j], height=Inches(6))
                                except:
                                    word.add_paragraph("You had not made a nomogram")
                            elif (j == 8):
                                word.add_paragraph("each-validation c-index:")
                                word.add_paragraph(str(self.operative_process[i][1][j]))
                            elif (j == 9):
                                word.add_paragraph("mean-validation c-index:")
                                word.add_paragraph(str(self.operative_process[i][1][j]))
                else:
                    for j in np.arange(0, len(self.operative_process[i][1])):
                        text2 = str(self.operative_process[i][1][j])
                        word.add_paragraph(text2)
        else:
            pass
        word.save(self.Result_save_path)
